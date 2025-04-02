"""
Complete solution for generating PDFs and DOCX files from markdown content
Requirements: reportlab, python-docx
Install with: pip install reportlab python-docx
"""

import os
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class PDFGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.setup_custom_styles()

    def setup_custom_styles(self):
        # Header style
        self.header_style = ParagraphStyle(
            'HeaderStyle',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1a4d7c'),
            spaceAfter=20
        )
        
        # Label style
        self.label_style = ParagraphStyle(
            'LabelStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#666666'),
            fontName='Helvetica-Bold'
        )
        
        # Value style
        self.value_style = ParagraphStyle(
            'ValueStyle',
            parent=self.styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#000000'),
            leading=14
        )
        
        # Description style
        self.description_style = ParagraphStyle(
            'DescriptionStyle',
            parent=self.styles['Normal'],
            fontSize=11,
            leading=14,
            textColor=colors.HexColor('#333333')
        )

    def create_product_pdf(self, content, filename):
        """Create a PDF product specification."""
        doc = SimpleDocTemplate(filename, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        story = []
        
        # Add header
        story.append(Paragraph("PRODUCT SPECIFICATION", self.header_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Parse content
        lines = content.strip().split('\n')
        data = {}
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                data[key.strip()] = value.strip()
        
        # Product name
        story.append(Paragraph(data.get('Product Name', ''), self.header_style))
        story.append(Spacer(1, 0.1 * inch))
        
        # Create product info table
        product_info = [
            ['Category', data.get('Category', '')],
            ['Price', data.get('Price', '')],
        ]
        
        t = Table(product_info, colWidths=[2*inch, 4*inch])
        t.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f5f5f5')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.2 * inch))
        
        # Description
        story.append(Paragraph("Description:", self.label_style))
        story.append(Paragraph(data.get('Description', ''), self.description_style))
        
        doc.build(story)

    def create_order_pdf(self, content, filename):
        """Create a PDF invoice for orders."""
        doc = SimpleDocTemplate(filename, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        story = []
        
        # Add header
        story.append(Paragraph("AVALANCHE WINTER GEAR", self.header_style))
        story.append(Paragraph("ORDER INVOICE", self.header_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Parse content into dictionary
        lines = content.strip().split('\n')
        data = {}
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                data[key.strip()] = value.strip()
        
        # Create order info table
        order_info = [
            [Paragraph("Order ID:", self.label_style), 
             Paragraph(data.get('Order ID', ''), self.value_style)],
            [Paragraph("Date:", self.label_style), 
             Paragraph(data.get('Date', ''), self.value_style)],
            [Paragraph("Customer ID:", self.label_style), 
             Paragraph(data.get('Customer ID', ''), self.value_style)]
        ]
        
        t = Table(order_info, colWidths=[1.5*inch, 4*inch])
        t.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.white),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f5f5f5')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.2 * inch))
        
        # Create product info table
        product_info = [
            ['Product Details', 'Price', 'Quantity', 'Total'],
            [data.get('Product Name', ''), 
             data.get('Price', ''),
             data.get('Quantity Ordered', ''),
             data.get('Total Price', '')]
        ]
        
        t = Table(product_info, colWidths=[3*inch, 1.2*inch, 1*inch, 1.2*inch])
        t.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a4d7c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(t)
        
        doc.build(story)

    def create_shipping_pdf(self, content, filename):
        """Create a PDF shipping receipt."""
        doc = SimpleDocTemplate(filename, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=72)
        
        story = []
        
        # Add header
        story.append(Paragraph("SHIPPING RECEIPT", self.header_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Parse content
        lines = content.strip().split('\n')
        data = {}
        for line in lines:
            if ':' in line:
                key, value = line.split(':', 1)
                data[key.strip()] = value.strip()
        
        # Create shipping info table
        shipping_info = [
            ['Tracking Information', ''],
            ['Order ID', data.get('Order ID', '')],
            ['Shipping Date', data.get('Shipping Date', '')],
            ['Carrier', data.get('Carrier', '')],
            ['Tracking Number', data.get('Tracking Number', '')],
            ['Status', data.get('Status', '')],
            ['Location', f"{data.get('Latitude', '')}, {data.get('Longitude', '')}"]
        ]
        
        t = Table(shipping_info, colWidths=[2*inch, 4*inch])
        t.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a4d7c')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('SPAN', (0, 0), (1, 0)),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#f5f5f5')),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
        ]))
        story.append(t)
        
        doc.build(story)

def split_customer_reviews(content, output_dir="customer-reviews"):
    """Split customer reviews into individual DOCX files."""
    os.makedirs(output_dir, exist_ok=True)
    
    reviews = content.strip().split("\n\n")
    
    for i, review in enumerate(reviews, 1):
        filename = os.path.join(output_dir, f"review-{i:02d}.docx")
        
        # Create a new document
        doc = Document()
        
        # Parse review content
        lines = review.strip().split('\n')
        product_name = ''
        review_text = ''
        review_date = ''
        
        for line in lines:
            if line.startswith('Product Name:'):
                product_name = line.split(':', 1)[1].strip()
            elif line.startswith('Review:'):
                review_text = line.split(':', 1)[1].strip().strip('"')
            elif line.startswith('Date:'):
                review_date = line.split(':', 1)[1].strip()
        
        # Add title
        title = doc.add_heading(level=1)
        title_run = title.add_run('Product Review')
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(26, 77, 124)  # Dark blue
        
        # Add product name
        product_para = doc.add_paragraph()
        product_run = product_para.add_run('Product: ')
        product_run.bold = True
        product_para.add_run(product_name)
        
        # Add date
        date_para = doc.add_paragraph()
        date_run = date_para.add_run('Date: ')
        date_run.bold = True
        date_para.add_run(review_date)
        
        # Add spacing
        doc.add_paragraph()
        
        # Add review header
        review_header = doc.add_paragraph()
        review_header_run = review_header.add_run('Customer Review')
        review_header_run.bold = True
        
        # Add review text
        review_para = doc.add_paragraph()
        review_para.add_run(review_text)
        
        # Save the document
        doc.save(filename)

def process_files():
    """Process all files and generate PDFs/DOCX files."""
    pdf_gen = PDFGenerator()
    
    try:
        # Create output directories with names matching MD files
        os.makedirs("product-catalog", exist_ok=True)
        os.makedirs("customer-reviews", exist_ok=True)
        os.makedirs("order-history", exist_ok=True)
        os.makedirs("shipping-logs", exist_ok=True)

        # Process Product Catalog (PDF)
        with open("product-catalog.md", 'r') as f:
            products = f.read().strip().split("\n\n")
            for i, product in enumerate(products, 1):
                filename = os.path.join("product-catalog", f"product-{i:02d}.pdf")
                pdf_gen.create_product_pdf(product, filename)
        print("Successfully processed Product Catalog")

        # Process Customer Reviews (DOCX)
        with open("customer-reviews.md", 'r') as f:
            content = f.read()
            split_customer_reviews(content)
        print("Successfully processed Customer Reviews")

        # Process Order History (PDF)
        with open("order-history.md", 'r') as f:
            orders = f.read().strip().split("\n\n")
            for i, order in enumerate(orders, 1):
                filename = os.path.join("order-history", f"order-{i:02d}.pdf")
                pdf_gen.create_order_pdf(order, filename)
        print("Successfully processed Order History")

        # Process Shipping Logs (PDF)
        with open("shipping-logs.md", 'r') as f:
            logs = f.read().strip().split("\n\n")
            for i, log in enumerate(logs, 1):
                filename = os.path.join("shipping-logs", f"shipping-{i:02d}.pdf")
                pdf_gen.create_shipping_pdf(log, filename)
        print("Successfully processed Shipping Logs")

    except FileNotFoundError as e:
        print(f"Error: File not found - {e.filename}")
    except Exception as e:
        print(f"Error processing files: {str(e)}")

if __name__ == "__main__":
    process_files()
